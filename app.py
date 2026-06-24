import streamlit as st
import google.generativeai as genai
import PyPDF2
import io
import re
from docx import Document
from fpdf import FPDF

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="NCDC English Exam Architect",
    page_icon="📝",
    layout="wide"
)

# --- CLEAN CHILD-FRIENDLY TYPOGRAPHY ---
st.markdown("""
    <style>
    * {
        font-family: 'Ubuntu', 'Segoe UI', 'Arial', sans-serif !important;
    }
    .blueprint-header {
        background-color: #f8f9fa;
        border-left: 5px solid #28a745;
        padding: 15px;
        border-radius: 4px;
        margin-bottom: 20px;
    }
    </style>
""", unsafe_allow_html=True)

# --- HELPER FUNCTIONS FOR DOCUMENT GENERATION ---
def create_word_docx(text, title):
    """Converts markdown text into a properly formatted Word Document"""
    doc = Document()
    doc.add_heading(title, 0)
    
    if not text:
        text = "No content generated."
        
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: 
                    run.bold = True
                    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_pdf(text, title):
    """Converts text into a clean, printable PDF with absolute crash protection"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    if not text:
        text = "No content generated."
    
    # 1. Unicode Sanitization
    replacements = {
        '‘': "'", '’': "'", '“': '"', '”': '"', '–': '-', '—': '-', '…': '...'
    }
    for search, replace in replacements.items():
        text = text.replace(search, replace)
        
    text = text.encode('latin-1', errors='replace').decode('latin-1')
    
    # 2. Aggressive Punctuation Shrinking
    # This forces ANY long string of characters to become a safe, short length
    text = re.sub(r'_{8,}', '________', text)   # Shrink underscores
    text = re.sub(r'-{8,}', '--------', text)   # Shrink dashes
    text = re.sub(r'\.{8,}', '........', text)  # Shrink dots
    text = re.sub(r'={8,}', '========', text)   # Shrink equal signs
    
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.ln(5)
    
    pdf.set_font("Helvetica", size=11)
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
            
        # 3. Force-break any remaining word longer than 40 characters
        words = line.split()
        safe_words = []
        for word in words:
            if len(word) > 40:
                safe_words.append(word[:40] + "-") # Hard crop with a dash
            else:
                safe_words.append(word)
        safe_line = " ".join(safe_words)
            
        if safe_line.startswith('### ') or safe_line.startswith('## ') or safe_line.startswith('# '):
            pdf.set_font("Helvetica", style="B", size=13)
            clean_heading = safe_line.replace('# ', '').replace('## ', '').replace('### ', '').replace('**', '')
            try:
                pdf.multi_cell(0, 8, clean_heading)
            except Exception:
                pass # Skip silently if a heading fails
            pdf.set_font("Helvetica", size=11)
        else:
            clean_line = safe_line.replace('**', '')
            # 4. Final Safety Net: If the PDF engine still panics, skip the line completely
            try:
                pdf.multi_cell(0, 6, clean_line)
            except Exception:
                pdf.cell(0, 6, "[Formatting error in line removed]", ln=True)
                
    pdf_out = pdf.output(dest='S')
    if isinstance(pdf_out, str):
        return pdf_out.encode('latin-1', errors='replace')
    return bytes(pdf_out)

# --- INITIALIZE SESSION STATE ---
if "generated_exam" not in st.session_state:
    st.session_state.generated_exam = None
if "generated_key" not in st.session_state:
    st.session_state.generated_key = None

# --- SIDEBAR: BLUEPRINT PARAMETERS ---
st.sidebar.header("📋 Target Class Configuration")
target_class = st.sidebar.selectbox("Select Target Class", ["Primary 5 (P.5)", "Primary 6 (P.6)", "Primary 7 (P.7)"])
target_term = st.sidebar.selectbox("Target Term", ["Term I", "Term II", "Term III", "Full Year Comprehensive"])
difficulty = st.sidebar.selectbox("Target Calibration", ["Standard UNEB Match", "Remedial Baseline", "Advanced Stretch"])

if st.sidebar.button("🗑️ Reset Application", use_container_width=True):
    st.session_state.generated_exam = None
    st.session_state.generated_key = None
    st.rerun()

# --- MAIN INTERFACE ---
st.title("📝 NCDC English Exam Architect")
st.subheader(f"Dynamic Assessment Generator — {target_class} ({target_term})")

# --- PDF FILE UPLOAD & TEXT EXTRACTION ---
st.markdown("### 📄 Step 1: Upload NCDC Curriculum Document")
uploaded_file = st.file_uploader("Upload your NCDC English Curriculum (.pdf)", type=["pdf"])

curriculum_text = ""
validation_passed = True

if uploaded_file is not None:
    with st.spinner("Extracting text from PDF pages..."):
        try:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    curriculum_text += extracted + "\n"
        except Exception as e:
            st.error(f"Error reading the PDF file: {str(e)}")
            validation_passed = False

    if curriculum_text and validation_passed:
        class_num = target_class.split(" ")[1] 
        detected_wrong_classes = [num for num in ["5", "6", "7"] if num != class_num and f"Primary {num}" in curriculum_text]
        
        if detected_wrong_classes and f"Primary {class_num}" not in curriculum_text:
            st.error(f"Curriculum Mismatch Warning: You selected {target_class}, but the PDF mentions Primary {', '.join(detected_wrong_classes)}.")
            validation_passed = False

# --- SYSTEM EXECUTION FORM ---
st.markdown("### 🤖 Step 2: Access Token & Run")
with st.form("generation_form"):
    api_key_input = st.text_input("Enter Gemini API Key", type="password")
    submit_btn = st.form_submit_button("🔨 Generate Examination & Answer Key", use_container_width=True)

# --- PROCESSING PIPELINE ---
if submit_btn:
    if not api_key_input or uploaded_file is None or not validation_passed:
        st.error("⚠️ Please resolve file uploads or API key issues before generating.")
    else:
        with st.spinner(f"Building Exam and formatting Answer Keys for {target_class}..."):
            
            system_instruction = (
                "You are an expert Primary School English item writer and curriculum analyst working with the NCDC. "
                "CRITICAL FORMATTING RULES:\n"
                "1. Output purely in clean Markdown text. No LaTeX.\n"
                "2. You MUST generate the Examination Paper first.\n"
                "3. You MUST separate the Examination Paper from the Marking Guide by using exactly this delimiter on its own line: === SPLIT HERE ===\n"
                "4. After the delimiter, provide the explicit Marking Guide/Answer Key.\n\n"
                "EXAM BLUEPRINT:\n"
                "SECTION A (50 Marks): Q1–5 Gap-filling, Q6–15 Morphology, Q16–30 Discrete Skills, Q31–50 Syntactic Transformations.\n"
                "SECTION B (50 Marks): Q51 Prose, Q52 Poetry, Q53 Functional Literacy, Q54 Narrative Discourse, Q55 Formal Composition."
            )
            
            user_prompt = f"""
            Target Class Level: {target_class}
            Target Academic Term: {target_term}
            Assessment Calibration Mode: {difficulty}
            --- CURRICULUM SOURCE ---
            {curriculum_text}
            --- END OF SOURCE ---
            
            Generate the 100-mark paper, write the delimiter "=== SPLIT HERE ===", and then generate the Answer Key.
            """
            
            try:
                genai.configure(api_key=api_key_input)
                model = genai.GenerativeModel(model_name="gemini-2.5-flash", system_instruction=system_instruction)
                response = model.generate_content(user_prompt)
                
                full_text = response.text
                if "=== SPLIT HERE ===" in full_text:
                    parts = full_text.split("=== SPLIT HERE ===")
                    st.session_state.generated_exam = parts[0].strip()
                    st.session_state.generated_key = parts[1].strip() if len(parts) > 1 else "Answer key missing."
                else:
                    st.session_state.generated_exam = full_text
                    st.session_state.generated_key = "Answer key was not separated. Please check the bottom of the exam document."
                
                st.success("✨ Files Successfully Generated and Formatted!")
            except Exception as e:
                st.error(f"Generation failed: {str(e)}")

# --- RENDER RESULTS & DOWNLOAD BUTTONS ---
if st.session_state.generated_exam:
    st.divider()
    
    st.markdown("### 📥 Download Your Documents")
    st.write("Your files have been automatically split and formatted for printing.")
    
    exam_pdf = create_pdf(st.session_state.generated_exam, f"{target_class} - English Examination")
    exam_word = create_word_docx(st.session_state.generated_exam, f"{target_class} - English Examination")
    key_word = create_word_docx(st.session_state.generated_key, f"{target_class} - Answer Key")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.download_button(
            label="📄 Download Exam (PDF)",
            data=exam_pdf,
            file_name=f"{target_class.replace(' ', '_')}_Exam.pdf",
            mime="application/pdf",
            use_container_width=True
        )
        
    with col2:
        st.download_button(
            label="📝 Download Exam (Word)",
            data=exam_word,
            file_name=f"{target_class.replace(' ', '_')}_Exam.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
    with col3:
        st.download_button(
            label="🔑 Download Answer Key (Word)",
            data=key_word,
            file_name=f"{target_class.replace(' ', '_')}_Answer_Key.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with st.expander("👁️ Preview Exam Text on Screen"):
        st.markdown(st.session_state.generated_exam)
    with st.expander("👁️ Preview Answer Key on Screen"):
        st.markdown(st.session_state.generated_key)