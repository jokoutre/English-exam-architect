import streamlit as st
import google.generativeai as genai
import PyPDF2
import io
import re
import textwrap
from docx import Document
from fpdf import FPDF

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="NCDC Science Exam Architect",
    page_icon="🔬",
    layout="wide"
)

# --- CLEAN SCIENCE TYPOGRAPHY ---
st.markdown("""
    <style>
    * {
        font-family: 'Ubuntu', 'Segoe UI', 'Arial', sans-serif !important;
    }
    .blueprint-header {
        background-color: #e8f5e9;
        border-left: 5px solid #2e7d32;
        padding: 15px;
        border-radius: 4px;
        margin-bottom: 20px;
    }
    </style>
""", unsafe_allow_html=True)

# --- BULLETPROOF DOCUMENT GENERATORS ---
def create_word_docx(text, title):
    doc = Document()
    doc.add_heading(title, 0)
    if not text: text = "No content generated."
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '): doc.add_heading(line.replace('# ', ''), level=1)
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_pdf(text, title):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    if not text: text = "No content generated."
    
    replacements = {'‘': "'", '’': "'", '“': '"', '”': '"', '–': '-', '—': '-', '…': '...'}
    for search, replace in replacements.items(): text = text.replace(search, replace)
    text = text.encode('latin-1', errors='replace').decode('latin-1')
    
    # The Sledgehammer Fix for long blank lines
    text = re.sub(r'_{8,}', '________', text)   
    text = re.sub(r'-{8,}', '--------', text)   
    text = re.sub(r'\.{8,}', '........', text)  
    text = re.sub(r'={8,}', '========', text)   
    
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.ln(5)
    
    pdf.set_font("Helvetica", size=11)
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
            
        words = line.split()
        safe_words = []
        for word in words:
            if len(word) > 40: safe_words.append(word[:40] + "-") 
            else: safe_words.append(word)
        safe_line = " ".join(safe_words)
            
        if safe_line.startswith('### ') or safe_line.startswith('## ') or safe_line.startswith('# '):
            pdf.set_font("Helvetica", style="B", size=13)
            clean_heading = safe_line.replace('# ', '').replace('## ', '').replace('### ', '').replace('**', '')
            try: pdf.multi_cell(0, 8, clean_heading)
            except: pass 
            pdf.set_font("Helvetica", size=11)
        else:
            clean_line = safe_line.replace('**', '')
            try: pdf.multi_cell(0, 6, clean_line)
            except: pdf.cell(0, 6, "[Formatting error in line removed]", ln=True)
                
    pdf_out = pdf.output(dest='S')
    if isinstance(pdf_out, str): return pdf_out.encode('latin-1', errors='replace')
    return bytes(pdf_out)

# --- INITIALIZE SESSION STATE ---
if "generated_sci" not in st.session_state: st.session_state.generated_sci = None
if "generated_sci_key" not in st.session_state: st.session_state.generated_sci_key = None

# --- SIDEBAR CONFIGURATION ---
st.sidebar.header("📋 Target Class Configuration")
# Primary 4 added successfully below!
target_class = st.sidebar.selectbox("Select Target Class", ["Primary 4 (P.4)", "Primary 5 (P.5)", "Primary 6 (P.6)", "Primary 7 (P.7)"])
target_term = st.sidebar.selectbox("Target Term", ["Term I", "Term II", "Term III", "UNEB Mock Standard"])

if st.sidebar.button("🗑️ Reset Application", use_container_width=True):
    st.session_state.generated_sci = None
    st.session_state.generated_sci_key = None
    st.rerun()

# --- MAIN INTERFACE ---
st.title("🔬 NCDC Science Exam Architect")
st.subheader(f"2025-Calibrated Integrated Science Assessment — {target_class}")

st.markdown("""
<div class="blueprint-header">
    <strong>Enforced 2025 UNEB Science Blueprint (100 Marks Total):</strong><br>
    • <b>Section A (40 Marks):</b> 40 short-answer questions. Includes visual interpretation placeholders and local scenario analysis.<br>
    • <b>Section B (60 Marks):</b> 15 multi-part questions (Q41-55). Strictly formatted with (a), (b), (c) and Roman numerals. Focuses heavily on Primary Health Care, agriculture, and applied physics.
</div>
""", unsafe_allow_html=True)

# --- PDF FILE UPLOAD ---
st.markdown("### 📄 Step 1: Upload NCDC Science Curriculum Document")
uploaded_file = st.file_uploader("Upload your NCDC Science Curriculum (.pdf)", type=["pdf"])

curriculum_text = ""
validation_passed = True

if uploaded_file is not None:
    with st.spinner("Extracting topics from PDF pages..."):
        try:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted: curriculum_text += extracted + "\n"
        except Exception as e:
            st.error(f"Error reading the PDF file: {str(e)}")
            validation_passed = False

    if curriculum_text and validation_passed:
        class_num = target_class.split(" ")[1] 
        # Added "4" to the safety check below!
        detected_wrong_classes = [num for num in ["4", "5", "6", "7"] if num != class_num and f"Primary {num}" in curriculum_text]
        if detected_wrong_classes and f"Primary {class_num}" not in curriculum_text:
            st.warning(f"Curriculum Warning: You selected {target_class}, but the PDF mentions Primary {', '.join(detected_wrong_classes)}.")

# --- SYSTEM EXECUTION FORM ---
st.markdown("### 🤖 Step 2: Access Token & Generate")
with st.form("generation_form"):
    api_key_input = st.text_input("Enter Gemini API Key", type="password")
    submit_btn = st.form_submit_button("🔨 Construct 100-Mark Science Paper", use_container_width=True)

# --- PROCESSING PIPELINE ---
if submit_btn:
    if not api_key_input or uploaded_file is None:
        st.error("⚠️ Please resolve file uploads or API key issues before generating.")
    else:
        with st.spinner(f"Building 2025-Calibrated Science exam for {target_class}..."):
            
            system_instruction = (
                "You are an expert Chief Examiner for UNEB setting the Primary Integrated Science Examination. "
                "CRITICAL DESIGN RULES:\n"
                "1. Output purely in clean Markdown text. NO LaTeX.\n"
                "2. The paper MUST highly resemble the actual 2025 UNEB format.\n"
                "3. Emphasize localized Ugandan scenarios: Primary Health Care (PHC), crop/animal husbandry (apiaries, goats, coffee), sanitation, and local environmental issues (charcoal burning).\n"
                "4. DIAGRAMS ARE MANDATORY: You must include at least 5 questions that rely on visual data. Use text placeholders like: '[Insert bar graph showing immunization of children here]' or '[Insert diagram of a root tuber here]' followed by analysis questions.\n"
                "5. You MUST separate the Examination Paper from the Marking Guide using exactly this delimiter on its own line: === SPLIT HERE ===\n\n"
                "SCIENCE EXAM BLUEPRINT:\n"
                "SECTION A (40 Marks): Exactly 40 short-answer questions (Questions 1-40). 1 mark each. Provide a blank line for the answer.\n"
                "SECTION B (60 Marks): Exactly 15 multi-part questions (Questions 41-55). 4 marks each. Structure these using (a), (b), (c) and Roman numerals (i), (ii) for listed answers. Include questions on mixed farming, human anatomy, simple machines, and disease control."
            )
            
            user_prompt = f"""
            Target Class Level: {target_class}
            Target Academic Term: {target_term}
            --- SCIENCE CURRICULUM SOURCE ---
            {curriculum_text}
            --- END OF SOURCE ---
            
            Generate the 100-mark science paper based strictly on the topics in the curriculum source and the 2025 blueprint rules. Write the delimiter "=== SPLIT HERE ===", and then generate the Answer Key.
            """
            
            try:
                genai.configure(api_key=api_key_input)
                model = genai.GenerativeModel(model_name="gemini-2.5-flash", system_instruction=system_instruction)
                response = model.generate_content(user_prompt)
                
                full_text = response.text
                if "=== SPLIT HERE ===" in full_text:
                    parts = full_text.split("=== SPLIT HERE ===")
                    st.session_state.generated_sci = parts[0].strip()
                    st.session_state.generated_sci_key = parts[1].strip() if len(parts) > 1 else "Answer key missing."
                else:
                    st.session_state.generated_sci = full_text
                    st.session_state.generated_sci_key = "Answer key was not separated."
                
                st.success("✨ Science Paper Successfully Generated!")
            except Exception as e:
                st.error(f"Generation failed: {str(e)}")

# --- RENDER RESULTS & DOWNLOAD BUTTONS ---
if st.session_state.generated_sci:
    st.divider()
    st.markdown("### 📥 Download Your Documents")
    
    exam_pdf = create_pdf(st.session_state.generated_sci, f"{target_class} - Science Examination")
    exam_word = create_word_docx(st.session_state.generated_sci, f"{target_class} - Science Examination")
    key_word = create_word_docx(st.session_state.generated_sci_key, f"{target_class} - Science Answer Key")
    
    col1, col2, col3 = st.columns(3)
    with col1: st.download_button("📄 Download Exam (PDF)", data=exam_pdf, file_name=f"{target_class.replace(' ', '_')}_Science.pdf", mime="application/pdf", use_container_width=True)
    with col2: st.download_button("📝 Download Exam (Word)", data=exam_word, file_name=f"{target_class.replace(' ', '_')}_Science.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    with col3: st.download_button("🔑 Download Answer Key (Word)", data=key_word, file_name=f"{target_class.replace(' ', '_')}_Science_Key.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

    with st.expander("👁️ Preview Science Exam"): st.markdown(st.session_state.generated_sci)
    with st.expander("👁️ Preview Science Answer Key"): st.markdown(st.session_state.generated_sci_key)