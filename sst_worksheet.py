import streamlit as st
import google.generativeai as genai
import PyPDF2
import io
import re
import textwrap
from docx import Document
from fpdf import FPDF

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="NCDC SST Worksheet Architect",
    page_icon="🌍",
    layout="wide"
)

# --- CLEAN WORKSHEET TYPOGRAPHY (Amber/Earth Theme) ---
st.markdown("""
    <style>
    * {
        font-family: 'Ubuntu', 'Segoe UI', 'Arial', sans-serif !important;
    }
    .blueprint-header {
        background-color: #fff8e1;
        border-left: 5px solid #ffb300;
        padding: 15px;
        border-radius: 4px;
        margin-bottom: 20px;
    }
    </style>
""", unsafe_allow_html=True)

# --- INVINCIBLE DOCUMENT GENERATORS ---
def create_word_docx(text, title):
    doc = Document()
    doc.add_heading(title, 0)
    if not text: text = "No content generated."
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '): doc.add_heading(line.replace('# ', ''), level=1)
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_pdf(text, title):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    if not text: text = "No content generated."
    
    replacements = {'‘': "'", '’': "'", '“': '"', '”': '"', '–': '-', '—': '-', '…': '...'}
    for search, replace in replacements.items(): text = text.replace(search, replace)
    text = text.encode('latin-1', errors='replace').decode('latin-1')
    
    # The Sledgehammer Fix for long blank lines / underscores
    text = re.sub(r'_{8,}', '________', text)   
    text = re.sub(r'-{8,}', '--------', text)   
    text = re.sub(r'\.{8,}', '........', text)  
    text = re.sub(r'={8,}', '========', text)   
    
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.ln(5)
    
    pdf.set_font("Helvetica", size=11)
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
            
        words = line.split()
        safe_words = []
        for word in words:
            if len(word) > 40: safe_words.append(word[:40] + "-") 
            else: safe_words.append(word)
        safe_line = " ".join(safe_words)
            
        if safe_line.startswith('### ') or safe_line.startswith('## ') or safe_line.startswith('# '):
            pdf.set_font("Helvetica", style="B", size=13)
            clean_heading = safe_line.replace('# ', '').replace('## ', '').replace('### ', '').replace('**', '')
            try: pdf.multi_cell(0, 8, clean_heading)
            except: pass 
            pdf.set_font("Helvetica", size=11)
        else:
            clean_line = safe_line.replace('**', '')
            try: pdf.multi_cell(0, 6, clean_line)
            except: pdf.cell(0, 6, "[Formatting error in line removed]", ln=True)
                
    pdf_out = pdf.output(dest='S')
    if isinstance(pdf_out, str): return pdf_out.encode('latin-1', errors='replace')
    return bytes(pdf_out)

# --- INITIALIZE SESSION STATE ---
if "generated_sst_ws" not in st.session_state: st.session_state.generated_sst_ws = None
if "generated_sst_ws_key" not in st.session_state: st.session_state.generated_sst_ws_key = None

# --- SIDEBAR CONFIGURATION ---
st.sidebar.header("📋 Target Parameter Configuration")
target_class = st.sidebar.selectbox("Select Target Class", ["Primary 4 (P.4)", "Primary 5 (P.5)", "Primary 6 (P.6)", "Primary 7 (P.7)"])
target_term = st.sidebar.selectbox("Target Term Scope", ["Term I", "Term II", "Term III"])

# THE MICRO-TARGETING FEATURE:
target_topic = st.sidebar.text_input("Worksheet Focus Topic / Theme", placeholder="e.g., Early Migrations, Climate, or Leaders in our District")

if st.sidebar.button("🗑️ Reset Application", use_container_width=True):
    st.session_state.generated_sst_ws = None
    st.session_state.generated_sst_ws_key = None
    st.rerun()

# --- MAIN INTERFACE ---
st.title("🌍 NCDC Daily SST Worksheet Architect")
st.subheader(f"Focused Topic Review — {target_class}")

st.markdown("""
<div class="blueprint-header">
    <strong>Enforced 10/10 Classroom Worksheet Blueprint:</strong><br>
    • <b>Section A: 10 Short Questions</b> Quick facts based on the target theme. <em>Question 10 is an EITHER/OR Religious Education question.</em><br>
    • <b>Section B: 10 Extended Tasks</b> Multi-part questions (a, b, c) and map/chart placeholders based on the theme. <em>Question 20 is a multi-part EITHER/OR Religious Education question.</em>
</div>
""", unsafe_allow_html=True)

# --- PDF FILE UPLOAD ---
st.markdown("### 📄 Step 1: Upload NCDC SST Curriculum Document")
uploaded_file = st.file_uploader("Upload your NCDC SST Curriculum (.pdf)", type=["pdf"])

curriculum_text = ""
validation_passed = True

if uploaded_file is not None:
    with st.spinner("Analyzing curriculum scope..."):
        try:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted: curriculum_text += extracted + "\n"
        except Exception as e:
            st.error(f"Error reading the PDF file: {str(e)}")
            validation_passed = False

# --- SYSTEM EXECUTION FORM ---
st.markdown("### 🤖 Step 2: Access Token & Generate Workspace")
with st.form("generation_form"):
    api_key_input = st.text_input("Enter Gemini API Key", type="password")
    submit_btn = st.form_submit_button("🔨 Construct 20-Question Targeted Worksheet", use_container_width=True)

# --- PROCESSING PIPELINE ---
if submit_btn:
    if not api_key_input or uploaded_file is None or not target_topic:
        st.error("⚠️ Missing Information: Please check your API key, PDF file upload, and make sure a Focus Topic is written in the sidebar.")
    else:
        with st.spinner(f"Drafting targeted worksheet for '{target_topic}'..."):
            
            system_instruction = (
                "You are an expert Primary SST curriculum designer creating classroom exercises based on the NCDC syllabus. "
                "CRITICAL INSTRUCTIONS:\n"
                "1. Output purely in clean Markdown text. NO LaTeX.\n"
                "2. The entire worksheet MUST focus primarily on the Target Focus Topic written by the teacher (e.g., 'Vegetation of Africa' or 'Transport and Communication').\n"
                "3. Use familiar local Ugandan or East African contexts.\n"
                "4. You MUST separate the Worksheet from the Answer Key using exactly this delimiter on its own line: === SPLIT HERE ===\n\n"
                "WORKSHEET STRUCTURE:\n"
                "SECTION A: Short-Answer Foundation (Questions 1-10). \n"
                "- Q1-Q9: Short recall questions based on the topic. Provide a short blank line `______` for answers.\n"
                "- Q10: MUST be an 'EITHER (Christian) OR (Islamic)' short-answer religious question, loosely tied to the moral aspects of the topic if possible.\n\n"
                "SECTION B: Application & Maps (Questions 11-20).\n"
                "- Q11-Q19: Multi-part questions broken down into (a), (b), (c). Include at least ONE placeholder for a diagram (e.g., '[Insert sketch map of Uganda showing... here]' or '[Insert pie chart here]') and ask questions about it.\n"
                "- Q20: MUST be a multi-part 'EITHER (Christian) OR (Islamic)' religious question."
            )
            
            user_prompt = f"""
            Target Class Level: {target_class}
            Target Scope: {target_term}
            Target Focus Topic: {target_topic}
            --- NCDC CURRICULUM CONTEXT ---
            {curriculum_text}
            --- END OF SOURCE ---
            
            Generate the targeted 20-question SST worksheet based on the focus topic and instructions (ensure the RE EITHER/OR sections are correct for Q10 and Q20), write the delimiter "=== SPLIT HERE ===", and then generate the Answer Key.
            """
            
            try:
                genai.configure(api_key=api_key_input)
                model = genai.GenerativeModel(model_name="gemini-2.5-flash", system_instruction=system_instruction)
                response = model.generate_content(user_prompt)
                
                full_text = response.text
                if "=== SPLIT HERE ===" in full_text:
                    parts = full_text.split("=== SPLIT HERE ===")
                    st.session_state.generated_sst_ws = parts[0].strip()
                    st.session_state.generated_sst_ws_key = parts[1].strip() if len(parts) > 1 else "Answer key missing."
                else:
                    st.session_state.generated_sst_ws = full_text
                    st.session_state.generated_sst_ws_key = "Answer key was not separated."
                
                st.success("✨ Targeted Classroom Worksheet Successfully Generated!")
            except Exception as e:
                st.error(f"Generation failed: {str(e)}")

# --- RENDER RESULTS & DOWNLOAD BUTTONS ---
if st.session_state.generated_sst_ws:
    st.divider()
    st.markdown("### 📥 Download Your Classroom Materials")
    
    clean_title = f"{target_class}_{target_topic.replace(' ', '_')}_Worksheet"
    
    exam_pdf = create_pdf(st.session_state.generated_sst_ws, f"{target_class} - SST Worksheet ({target_topic})")
    exam_word = create_word_docx(st.session_state.generated_sst_ws, f"{target_class} - SST Worksheet ({target_topic})")
    key_word = create_word_docx(st.session_state.generated_sst_ws_key, f"{target_class} - SST Worksheet Answer Key")
    
    col1, col2, col3 = st.columns(3)
    with col1: st.download_button("📄 Download Worksheet (PDF)", data=exam_pdf, file_name=f"{clean_title}.pdf", mime="application/pdf", use_container_width=True)
    with col2: st.download_button("📝 Download Worksheet (Word)", data=exam_word, file_name=f"{clean_title}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    with col3: st.download_button("🔑 Download Answer Key (Word)", data=key_word, file_name=f"{clean_title}_Key.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

    with st.expander("👁️ Preview Worksheet"): st.markdown(st.session_state.generated_sst_ws)
    with st.expander("👁️ Preview Answer Key"): st.markdown(st.session_state.generated_sst_ws_key)