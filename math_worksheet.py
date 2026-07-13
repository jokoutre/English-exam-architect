import streamlit as st
import google.generativeai as genai
import PyPDF2
import io
import re
import textwrap
from docx import Document
from fpdf import FPDF

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="NCDC Math Worksheet Architect",
    page_icon="📐",
    layout="wide"
)

# --- CLEAN WORKSHEET TYPOGRAPHY (Soft Blue Workspace Theme) ---
st.markdown("""
    <style>
    * {
        font-family: 'Ubuntu', 'Segoe UI', 'Arial', sans-serif !important;
    }
    .blueprint-header {
        background-color: #e3f2fd;
        border-left: 5px solid #0288d1;
        padding: 15px;
        border-radius: 4px;
        margin-bottom: 20px;
    }
    </style>
""", unsafe_allow_html=True)

# --- INVINCIBLE DOCUMENT GENERATORS ---
def create_word_docx(text, title):
    doc = Document()
    doc.add_heading(title, 0)
    if not text: text = "No content generated."
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '): doc.add_heading(line.replace('# ', ''), level=1)
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_pdf(text, title):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    if not text: text = "No content generated."
    
    replacements = {'‘': "'", '’': "'", '“': '"', '”': '"', '–': '-', '—': '-', '…': '...'}
    for search, replace in replacements.items(): text = text.replace(search, replace)
    text = text.encode('latin-1', errors='replace').decode('latin-1')
    
    # The Sledgehammer Fix for long blank lines / underscores
    text = re.sub(r'_{8,}', '________', text)   
    text = re.sub(r'-{8,}', '--------', text)   
    text = re.sub(r'\.{8,}', '........', text)  
    text = re.sub(r'={8,}', '========', text)   
    
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.ln(5)
    
    pdf.set_font("Helvetica", size=11)
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
            
        words = line.split()
        safe_words = []
        for word in words:
            if len(word) > 40: safe_words.append(word[:40] + "-") 
            else: safe_words.append(word)
        safe_line = " ".join(safe_words)
            
        if safe_line.startswith('### ') or safe_line.startswith('## ') or safe_line.startswith('# '):
            pdf.set_font("Helvetica", style="B", size=13)
            clean_heading = safe_line.replace('# ', '').replace('## ', '').replace('### ', '').replace('**', '')
            try: pdf.multi_cell(0, 8, clean_heading)
            except: pass 
            pdf.set_font("Helvetica", size=11)
        else:
            clean_line = safe_line.replace('**', '')
            try: pdf.multi_cell(0, 6, clean_line)
            except: pdf.cell(0, 6, "[Formatting error in line removed]", ln=True)
                
    pdf_out = pdf.output(dest='S')
    if isinstance(pdf_out, str): return pdf_out.encode('latin-1', errors='replace')
    return bytes(pdf_out)

# --- INITIALIZE SESSION STATE ---
if "generated_ws" not in st.session_state: st.session_state.generated_ws = None
if "generated_ws_key" not in st.session_state: st.session_state.generated_ws_key = None

# --- SIDEBAR CONFIGURATION ---
st.sidebar.header("📋 Target Parameter Configuration")
target_class = st.sidebar.selectbox("Select Target Class", ["Primary 4 (P.4)", "Primary 5 (P.5)", "Primary 6 (P.6)", "Primary 7 (P.7)"])
target_term = st.sidebar.selectbox("Target Term Scope", ["Term I", "Term II", "Term III"])

# THE MICRO-TARGETING FEATURE:
target_topic = st.sidebar.text_input("Worksheet Focus Topic / Theme", placeholder="e.g., Fractions and Decimals")

if st.sidebar.button("🗑️ Reset Application", use_container_width=True):
    st.session_state.generated_ws = None
    st.session_state.generated_ws_key = None
    st.rerun()

# --- MAIN INTERFACE ---
st.title("📐 NCDC Daily Math Worksheet Architect")
st.subheader(f"Focused Topic Review — {target_class}")

st.markdown("""
<div class="blueprint-header">
    <strong>Enforced 10/10 Classroom Worksheet Blueprint:</strong><br>
    • <b>Section A: 10 Short Questions</b> Focuses entirely on quick mental recall, foundational skills, and direct calculation blocks.<br>
    • <b>Section B: 10 Extended Questions</b> Focused entirely on multi-step reasoning, word problems inside local Ugandan settings, and graphics/diagram prompts based on the specific target theme.
</div>
""", unsafe_allow_html=True)

# --- PDF FILE UPLOAD ---
st.markdown("### 📄 Step 1: Upload NCDC Mathematics Curriculum Document")
uploaded_file = st.file_uploader("Upload your NCDC Math Curriculum (.pdf)", type=["pdf"])

curriculum_text = ""
validation_passed = True

if uploaded_file is not None:
    with st.spinner("Analyzing curriculum scope..."):
        try:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted: curriculum_text += extracted + "\n"
        except Exception as e:
            st.error(f"Error reading the PDF file: {str(e)}")
            validation_passed = False

# --- SYSTEM EXECUTION FORM ---
st.markdown("### 🤖 Step 2: Access Token & Generate Workspace")
with st.form("generation_form"):
    api_key_input = st.text_input("Enter Gemini API Key", type="password")
    submit_btn = st.form_submit_button("🔨 Construct 20-Question Targeted Worksheet", use_container_width=True)

# --- PROCESSING PIPELINE ---
if submit_btn:
    if not api_key_input or uploaded_file is None or not target_topic:
        st.error("⚠️ Missing Information: Please check your API key, PDF file upload, and make sure a Focus Topic is written in the sidebar.")
    else:
        with st.spinner(f"Drafting targeted worksheet for '{target_topic}'..."):
            
            system_instruction = (
                "You are an expert Primary Mathematics curriculum designer creating classroom exercises. "
                "CRITICAL INSTRUCTIONS:\n"
                "1. Output purely in clean Markdown text. ABSOLUTELY NO LaTeX FORMATTING allowed (never use $, $$, or math blocks). Use regular text formatting like 1/2, 3/4, x, =, or the word 'degrees'.\n"
                "2. The entire worksheet MUST focus primarily on the Target Focus Topic written by the teacher, calibrated to the difficulty scope of the selected class.\n"
                "3. Use familiar local names and objects (Ugandan shilling in whole numbers, water tanks, jerrycans, markets).\n"
                "4. You MUST separate the Worksheet from the Answer Key using exactly this delimiter on its own line: === SPLIT HERE ===\n\n"
                "WORKSHEET STRUCTURE:\n"
                "SECTION A: Short-Answer Foundation (Questions 1-10). Provide 10 simple calculation or direct single-step questions matching the topic. Leave a short blank line for the pupil's answer.\n"
                "SECTION B: Deep Application & Reasoning (Questions 11-20). Provide 10 rich multi-step word problems, table analysis tasks, or graphic placeholders (e.g., 'Shade the diagram below to match...', '[Insert coordinate path prompt]') focused entirely on the topic."
            )
            
            user_prompt = f"""
            Target Class Level: {target_class}
            Target Scope: {target_term}
            Target Focus Topic: {target_topic}
            --- NCDC CURRICULUM CONTEXT ---
            {curriculum_text}
            --- END OF SOURCE ---
            
            Generate the targeted 20-question worksheet based on the focus topic and instructions, write the delimiter "=== SPLIT HERE ===", and then generate a detailed step-by-step Answer Key showing the calculations.
            """
            
            try:
                genai.configure(api_key=api_key_input)
                model = genai.GenerativeModel(model_name="gemini-2.5-flash", system_instruction=system_instruction)
                response = model.generate_content(user_prompt)
                
                full_text = response.text
                if "=== SPLIT HERE ===" in full_text:
                    parts = full_text.split("=== SPLIT HERE ===")
                    st.session_state.generated_ws = parts[0].strip()
                    st.session_state.generated_ws_key = parts[1].strip() if len(parts) > 1 else "Answer key missing."
                else:
                    st.session_state.generated_ws = full_text
                    st.session_state.generated_ws_key = "Answer key was not separated."
                
                st.success("✨ Targeted Classroom Worksheet Successfully Generated!")
            except Exception as e:
                st.error(f"Generation failed: {str(e)}")

# --- RENDER RESULTS & DOWNLOAD BUTTONS ---
if st.session_state.generated_ws:
    st.divider()
    st.markdown("### 📥 Download Your Classroom Materials")
    
    clean_title = f"{target_class}_{target_topic.replace(' ', '_')}_Worksheet"
    
    exam_pdf = create_pdf(st.session_state.generated_ws, f"{target_class} - Math Worksheet ({target_topic})")
    exam_word = create_word_docx(st.session_state.generated_ws, f"{target_class} - Math Worksheet ({target_topic})")
    key_word = create_word_docx(st.session_state.generated_ws_key, f"{target_class} - Math Worksheet Answer Key")
    
    col1, col2, col3 = st.columns(3)
    with col1: st.download_button("📄 Download Worksheet (PDF)", data=exam_pdf, file_name=f"{clean_title}.pdf", mime="application/pdf", use_container_width=True)
    with col2: st.download_button("📝 Download Worksheet (Word)", data=exam_word, file_name=f"{clean_title}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    with col3: st.download_button("🔑 Download Answer Key (Word)", data=key_word, file_name=f"{clean_title}_Key.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

    with st.expander("👁️ Preview Worksheet"): st.markdown(st.session_state.generated_ws)
    with st.expander("👁️ Preview Step-by-Step Answer Key"): st.markdown(st.session_state.generated_ws_key)